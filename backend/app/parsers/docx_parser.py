"""
DOCX parser — paragraph extraction via python-docx.

DOCX files have no native page number concept, so we group
paragraphs into synthetic "sections" and use a section index
as the page_number equivalent.
"""

from __future__ import annotations

from docx import Document as DocxDocument

from app.parsers.base import ParsedDocument

# Number of non-empty paragraphs to group into one section.
_PARAGRAPHS_PER_SECTION = 20


def parse_docx(file_bytes: bytes, filename: str) -> list[ParsedDocument]:
    """
    Extract text from a DOCX file grouped into sections.

    Parameters
    ----------
    file_bytes : raw bytes of the uploaded DOCX
    filename   : original filename (for metadata)

    Returns
    -------
    List of ParsedDocument, one per section.
    """
    import io

    doc = DocxDocument(io.BytesIO(file_bytes))

    # Collect all non-empty paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    if not paragraphs:
        return []

    sections: list[ParsedDocument] = []
    section_index = 1

    for i in range(0, len(paragraphs), _PARAGRAPHS_PER_SECTION):
        chunk = paragraphs[i : i + _PARAGRAPHS_PER_SECTION]
        text = "\n\n".join(chunk)

        sections.append(
            ParsedDocument(
                text=text,
                page_number=section_index,  # synthetic section index
                source_file=filename,
            )
        )
        section_index += 1

    return sections
